import asyncio
import io
import logging
import zipfile
from concurrent.futures import ThreadPoolExecutor

from lxml import etree
import docx
import mammoth

from app.extractors.base import BaseExtractor

logger = logging.getLogger(__name__)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

class DocxExtractor(BaseExtractor):
    """Extract text from .docx documents.

    Primary: python-docx for structured extraction with heading preservation.
    Fallback: mammoth raw text extraction.
    """

    def __init__(self, executor: ThreadPoolExecutor) -> None:
        self._executor = executor

    async def extract(self, content: bytes, filename: str = "") -> str:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self._executor, self._sync_extract, content, filename)

    def _sync_extract(self, content: bytes, filename: str) -> str:
        try:
            return self._extract_with_python_docx(content)
        except Exception as exc:
            logger.warning(
                "python-docx failed for %s (%s), falling back to mammoth",
                filename,
                type(exc).__name__,
            )
            return self._extract_with_mammoth(content)

    def _extract_with_python_docx(self, content: bytes) -> str:
        doc = docx.Document(io.BytesIO(content))
        parts: list[str] = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Find the matching Paragraph object
                for para in doc.paragraphs:
                    if para._element is element:
                        parts.append(self._format_paragraph(para))
                        break
            elif tag == "tbl":
                for table in doc.tables:
                    if table._element is element:
                        parts.append(self._format_table(table))
                        break

        body_text = "\n\n".join(p for p in parts if p.strip())

        notes_section = self._extract_notes(content)
        if notes_section:
            body_text = body_text + notes_section

        return body_text

    @staticmethod
    def _format_paragraph(para) -> str:
        style_name = (para.style.name or "").lower() if para.style else ""
        text = para.text.strip()
        if not text:
            return ""

        heading_map = {
            "heading 1": "# ",
            "heading 2": "## ",
            "heading 3": "### ",
            "heading 4": "#### ",
            "heading 5": "##### ",
            "heading 6": "###### ",
            "title": "# ",
            "subtitle": "## ",
        }

        prefix = heading_map.get(style_name, "")
        return f"{prefix}{text}"

    @staticmethod
    def _format_table(table) -> str:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        lines: list[str] = []
        # Header row
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
        # Data rows
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")

        return "\n".join(lines)

    def _extract_notes(self, content: bytes) -> str:
        """
        Extract footnotes and endnotes from a .docx zip.
        Returns a formatted string section, or empty string if none found.
        """
        sections = []

        for xml_path, section_title in [
            ("word/footnotes.xml", "Footnotes"),
            ("word/endnotes.xml", "Endnotes"),
        ]:
            notes = self._parse_notes_xml(content, xml_path)
            if notes:
                lines = [f"\n\n## {section_title}\n"]
                for i, text in enumerate(notes, start=1):
                    lines.append(f"[{i}] {text}")
                sections.append("\n".join(lines))

        return "\n".join(sections)

    def _parse_notes_xml(self, content: bytes, xml_path: str) -> list[str]:
        """
        Parse word/footnotes.xml or word/endnotes.xml and return a list of
        note text strings (skipping Word's internal separator entries).
        Returns empty list if the file does not exist in the zip.
        """
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as z:
                if xml_path not in z.namelist():
                    return []
                xml_bytes = z.read(xml_path)
        except (zipfile.BadZipFile, KeyError):
            return []

        try:
            root = etree.fromstring(xml_bytes)
        except etree.XMLSyntaxError:
            logger.warning("Malformed XML found in %s", xml_path)
            return []

        note_tag = f"{{{W_NS}}}footnote" if "footnotes" in xml_path else f"{{{W_NS}}}endnote"
        id_attr = f"{{{W_NS}}}id"
        p_tag = f"{{{W_NS}}}p"
        r_tag = f"{{{W_NS}}}r"
        t_tag = f"{{{W_NS}}}t"
        rstyle_tag = f"{{{W_NS}}}rStyle"

        # IDs -1 and 0 are Word's internal separator/continuation-separator entries
        SKIP_IDS = {"-1", "0"}

        results = []

        for note_elem in root.iter(note_tag):
            note_id = note_elem.get(id_attr, "")
            if note_id in SKIP_IDS:
                continue

            paragraphs = []
            for p_elem in note_elem.iter(p_tag):
                p_parts = []
                for r_elem in p_elem.iter(r_tag):
                    # Skip runs that are styled as FootnoteReference / EndnoteReference
                    is_ref = False
                    for rstyle in r_elem.iter(rstyle_tag):
                        style_val = rstyle.get(f"{{{W_NS}}}val", "")
                        if style_val in ("FootnoteReference", "EndnoteReference"):
                            is_ref = True
                            break
                    if is_ref:
                        continue

                    # Extract all text fragments within this valid run
                    for t_elem in r_elem.iter(t_tag):
                        if t_elem.text:
                            p_parts.append(t_elem.text)

                p_text = "".join(p_parts).strip()
                if p_text:
                    paragraphs.append(p_text)

            note_text = "\n".join(paragraphs).strip()
            if note_text:
                results.append(note_text)

        return results

    @staticmethod
    def _extract_with_mammoth(content: bytes) -> str:
        result = mammoth.extract_raw_text(io.BytesIO(content))
        return result.value
